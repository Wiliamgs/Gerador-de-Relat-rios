import os
import re
import pandas as pd
import PyPDF2
from flask import Flask, request, render_template, send_file
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Configuração do Flask ---
app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# --- MAPA DE PROFISSIONAIS ---
PROFESSIONAL_MAP = {
    "Amira Antonella Pontes de Almeida": {
        'Psicologia': 'Leonardo Santana Honorato', 'Psicomotricidade': 'Camila Barbosa dos Santos Silveiro',
        'Terapia Ocupacional': 'Laila Carolina Matos Serrão', 'Fonoaudiologia': 'João Vitor Ribeiro de Oliveira',
        'Psicopedagogia': 'Thais Cristina Matteucci'
    },
    "Arthur Pinheiro Bertoni": {
        'Psicologia': 'Sabrina de Souza Silva', 'Fisioterapia': 'Thais Gonçalves',
        'Terapia Ocupacional': 'Ana Carolina dos Santos Sarmento', 'Psicopedagogia': 'Thais Cristina Matteucci',
        'Fonoaudiologia': 'Dominique Pereira Santos Lanera'
    },
    "Augusto Quina Severo": {
        'Psicologia': 'Sabrina de Souza Silva', 'Fisioterapia': 'Thais Gonçalves',
        'Psicopedagogia': 'Larissa Bueno Ferreira', 'Terapia Ocupacional': 'Ana Carolina dos Santos Sarmento',
        'Fonoaudiologia': 'Dominique Pereira Santos Lanera'
    },
    "Gabriel Henrique Rocha Oliveira": {
        'Terapia Ocupacional': 'Laila Carolina Matos Serrão', 'Psicomotricidade': 'Camila Silvério',
        'Fisioterapia': 'Gabriel Almeida Gonçalves', 'Psicopedagogia': 'Larissa Bueno Ferreira',
        'Psicologia': 'Leonardo Santana Honorato', 'TO': 'Laila Carolina Matos Serrão'
    },
    "Heitor Gabriel Campos Conceição": {
        'Terapia Ocupacional': 'Laila Matos', 'Fonoaudiologia': 'João Vítor Ribeiro de Oliveira',
        'Psicologia': 'Leonardo Santana Honorato', 'Psicopedagogia': 'Christine Gomes',
        'Psicomotricidade': 'Camila Silvério', 'Fisioterapia': 'Thais Gonçalves'
    },
    "Jade Bandeira de Oliveira": {
        'Psicologia': 'Gabriela Ferreira da Conceição', 'Psicopedagogia': 'Thais Cristina Matteucci',
        'Fonoaudiologia': 'João Vitor Ribeiro de Oliveira', 'Terapia Ocupacional': 'Laila Matos',
        'Psicomotricidade': 'Camila Barbosa dos Santos Silveiro', 'Fisioterapia': 'Thais Gonçalves'
    },
    "João Pedro Roucourt Salviano": {
        'Terapia Ocupacional': 'Riverson Ronald Silva da Costa', 'Psicologia': 'Sabrina de Souza Silva',
        'Psicopedagogia': 'Larissa Bueno', 'Fonoaudiologia': 'Selma Maria Oliveira de Souza',
        'Musicoterapia': 'João Paulo Silva Lopes', 'Psicomotricidade': 'Camila Barbosa dos Santos Silveiro'
    },
    "Malu Alves Bicalho": {
        'Fonoaudiologia': 'João Vitor Ribeiro de Oliveira', 'Terapia Ocupacional': 'Laila Carolina Matos Serrão',
        'Psicopedagogia': 'Larissa Bueno Ferreira', 'Fisioterapia': 'Luana da Silva Santos',
        'Psicologia': 'Caroline Ribeiro Souza'
    },
    "Joseph Miguel Ferreira Diniz": {
        'Fisioterapia': 'A Definir', 'Terapia Ocupacional': 'A Definir', 
        'Fonoaudiologia': 'A Definir', 'CME': 'A Definir', 'Psicologia': 'A Definir'
    }
}

# --- Palavras Ignoradas (Feriados, Faltas e Atestados) ---
PALAVRAS_IGNORADAS = [
    'natal', 'ano novo', 'páscoa', 'carnaval', 'dia das crianças', 
    'dia das mães', 'dia dos pais', 'festa', 'festividade', 
    'feriado', 'férias', 'papai noel', 'comemoração', 'recesso',
    'falta', 'faltou', 'ausência', 'ausente', 'não compareceu', 
    'atestado', 'cancelado', 'cancelada', 'desmarcou', 'desmarcado'
]

# --- Funções de Processamento ---
def extrair_texto_de_pdf(caminho_pdf):
    texto_completo = ""
    with open(caminho_pdf, 'rb') as arquivo_pdf:
        leitor_pdf = PyPDF2.PdfReader(arquivo_pdf)
        for pagina in leitor_pdf.pages:
            try:
                texto_completo += pagina.extract_text() or ""
            except Exception as e:
                print(f"Erro ao extrair texto da página: {e}")
                continue
    return texto_completo

def get_schedule_for_patient(texto_pdf):
    paciente_identificado = None
    texto_pdf_normalizado = texto_pdf.lower()
    for nome_paciente in PROFESSIONAL_MAP.keys():
        if nome_paciente.lower() in texto_pdf_normalizado:
            paciente_identificado = nome_paciente
            break
            
    if not paciente_identificado:
        return None, None
    
    # --- LÓGICA MARÇO 2026 ---
    dados = []

    if paciente_identificado == "Amira Antonella Pontes de Almeida":
        for dia in ['04/03/2026', '05/03/2026', '18/03/2026']:
            dados.extend([(dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicologia'), (dia, 'Psicomotricidade')])
        for dia in ['06/03/2026', '20/03/2026', '27/03/2026']:
            dados.extend([(dia, 'Fonoaudiologia'), (dia, 'Psicologia'), (dia, 'Psicopedagogia')])
        for dia in ['11/03/2026', '25/03/2026']:
             dados.extend([(dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicologia')])
        for dia in ['19/03/2026', '26/03/2026']:
            dados.extend([(dia, 'Terapia Ocupacional'), (dia, 'Psicologia'), (dia, 'Psicomotricidade')])
        dados.extend([('24/03/2026', 'Terapia Ocupacional'), ('24/03/2026', 'Fonoaudiologia'), ('24/03/2026', 'Psicomotricidade')])

    elif paciente_identificado == "Arthur Pinheiro Bertoni":
        for dia in ['04/03/2026', '18/03/2026', '25/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Fisioterapia'), (dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicopedagogia')])

    elif paciente_identificado == "Augusto Quina Severo":
        for dia in ['03/03/2026', '17/03/2026', '24/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Psicopedagogia')])
        for dia in ['05/03/2026', '12/03/2026', '19/03/2026', '26/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Fisioterapia'), (dia, 'Psicologia')])
        dados.extend([('31/03/2026', 'Fisioterapia'), ('31/03/2026', 'Fisioterapia'), ('31/03/2026', 'Psicopedagogia')])

    elif paciente_identificado == "Gabriel Henrique Rocha Oliveira":
        # Paciente apenas com faltas em março, retorna lista vazia
        pass

    elif paciente_identificado == "Heitor Gabriel Campos Conceição":
        for dia in ['02/03/2026', '04/03/2026', '09/03/2026', '11/03/2026', '16/03/2026', '18/03/2026', '30/03/2026']:
            dados.extend([(dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicologia'), (dia, 'Psicopedagogia')])
        for dia in ['03/03/2026', '24/03/2026', '31/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicopedagogia')])
        for dia in ['05/03/2026', '12/03/2026', '13/03/2026', '19/03/2026', '20/03/2026', '27/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Fonoaudiologia'), (dia, 'Psicologia')])
        for dia in ['10/03/2026', '17/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia')])
        dados.extend([('23/03/2026', 'Fonoaudiologia'), ('23/03/2026', 'Psicologia'), ('23/03/2026', 'Psicopedagogia')])
        dados.extend([('26/03/2026', 'Fisioterapia'), ('26/03/2026', 'Fonoaudiologia'), ('26/03/2026', 'Psicologia'), ('26/03/2026', 'Psicomotricidade')])

    elif paciente_identificado == "Jade Bandeira de Oliveira":
        for dia in ['02/03/2026', '04/03/2026', '09/03/2026', '11/03/2026', '18/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Fonoaudiologia'), (dia, 'Psicopedagogia')])
        for dia in ['05/03/2026', '19/03/2026', '26/03/2026']:
            dados.extend([(dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicologia')])
        for dia in ['10/03/2026', '31/03/2026']:
             dados.extend([(dia, 'Fisioterapia'), (dia, 'Fonoaudiologia'), (dia, 'Psicologia')])
        dados.extend([('13/03/2026', 'Fisioterapia'), ('13/03/2026', 'Terapia Ocupacional')])
        dados.extend([('16/03/2026', 'Fisioterapia'), ('16/03/2026', 'Fonoaudiologia')])
        for dia in ['20/03/2026', '27/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia')])
        for dia in ['23/03/2026', '24/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicopedagogia')])
        dados.extend([('25/03/2026', 'Terapia Ocupacional'), ('25/03/2026', 'Fonoaudiologia'), ('25/03/2026', 'Psicopedagogia')])

    elif paciente_identificado == "João Pedro Roucourt Salviano":
        for dia in ['02/03/2026', '03/03/2026']:
            dados.extend([(dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicologia'), (dia, 'Psicopedagogia'), (dia, 'Psicomotricidade')])
        for dia in ['05/03/2026', '12/03/2026', '19/03/2026']:
             dados.extend([(dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicologia'), (dia, 'Psicopedagogia')])
        for dia in ['09/03/2026', '10/03/2026', '16/03/2026', '17/03/2026', '23/03/2026', '24/03/2026', '30/03/2026', '31/03/2026']:
             dados.extend([(dia, 'Terapia Ocupacional'), (dia, 'Psicologia'), (dia, 'Psicopedagogia')])

    elif paciente_identificado == "Malu Alves Bicalho":
        for dia in ['04/03/2026', '05/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia'), (dia, 'Psicologia')])
        for dia in ['09/03/2026', '16/03/2026', '23/03/2026', '30/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia')])
        dados.extend([('12/03/2026', 'Fisioterapia'), ('12/03/2026', 'Terapia Ocupacional'), ('12/03/2026', 'Fonoaudiologia'), ('12/03/2026', 'Psicologia'), ('12/03/2026', 'Psicopedagogia')])
        dados.extend([('13/03/2026', 'Fonoaudiologia'), ('13/03/2026', 'Psicologia')])
        for dia in ['17/03/2026', '24/03/2026']:
            dados.extend([(dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia')])
        for dia in ['18/03/2026', '25/03/2026']:
             dados.extend([(dia, 'Psicologia')])
        dados.extend([('20/03/2026', 'Fisioterapia'), ('20/03/2026', 'Fonoaudiologia'), ('20/03/2026', 'Psicologia')])
        dados.extend([('26/03/2026', 'Fonoaudiologia'), ('26/03/2026', 'Psicologia')])
        dados.extend([('27/03/2026', 'Fisioterapia')])

    elif paciente_identificado == "Joseph Miguel Ferreira Diniz":
        for dia in ['04/03/2026', '12/03/2026', '19/03/2026', '20/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Fisioterapia'), (dia, 'Fonoaudiologia')])
        for dia in ['05/03/2026', '06/03/2026', '09/03/2026', '10/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Fisioterapia'), (dia, 'Terapia Ocupacional'), (dia, 'Fonoaudiologia')])
        for dia in ['11/03/2026', '18/03/2026', '25/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Fisioterapia'), (dia, 'Fonoaudiologia'), (dia, 'CME')])
        for dia in ['16/03/2026', '17/03/2026', '24/03/2026', '31/03/2026']:
            dados.extend([(dia, 'Fisioterapia'), (dia, 'Fisioterapia'), (dia, 'Terapia Ocupacional'), (dia, 'CME')])
        dados.extend([('23/03/2026', 'Fisioterapia'), ('23/03/2026', 'Fisioterapia'), ('23/03/2026', 'Terapia Ocupacional'), ('23/03/2026', 'Fonoaudiologia'), ('23/03/2026', 'CME')])
        dados.extend([('26/03/2026', 'Fisioterapia'), ('26/03/2026', 'Fisioterapia'), ('26/03/2026', 'Terapia Ocupacional')])
        dados.extend([('30/03/2026', 'Fisioterapia'), ('30/03/2026', 'Fisioterapia')])

    return pd.DataFrame(dados, columns=['Data', 'Especialidade']), paciente_identificado

def analisar_prontuario(texto_pdf):
    evolucoes = []
    blocos = re.split(r'\nProfissional: ', '\n' + texto_pdf)[1:]
    stop_phrases = ["Av. Dr. Renato de Andrade Maia"]
    for bloco in blocos:
        try:
            profissional_completo = bloco.split('\n')[0].strip()
            match_especialidade = re.search(r"Especialidade: (.*?)\n", bloco)
            match_data = re.search(r"Data: (\d{2}/\d{2}/\d{4})", bloco)
            match_anotacoes = re.search(r"Anotações:\n(.*?)(?=\Z|Profissional:)", bloco, re.DOTALL)
            if match_data and match_especialidade and match_anotacoes:
                especialidade = match_especialidade.group(1).strip()
                data = match_data.group(1).strip()
                anotacoes = match_anotacoes.group(1).strip()
                anotacoes = re.sub(r'Inserido por.*\n?', '', anotacoes).strip()
                primeiro_nome = profissional_completo.split(' ')[0]
                if primeiro_nome and primeiro_nome in anotacoes:
                    anotacoes = anotacoes.split(primeiro_nome)[0].strip()
                for phrase in stop_phrases:
                    if phrase in anotacoes:
                        anotacoes = anotacoes.split(phrase)[0].strip()
                anotacoes = re.sub(r'\n\s*\n', '\n', anotacoes)
                if anotacoes:
                    evolucoes.append({'Profissional': profissional_completo, 'Especialidade': especialidade, 'Data': data, 'Anotacoes_prontuario': anotacoes})
        except Exception as e:
            continue
    return pd.DataFrame(evolucoes, columns=['Profissional', 'Especialidade', 'Data', 'Anotacoes_prontuario'])

def deve_ignorar_anotacao(texto):
    texto_lower = texto.lower()
    for palavra in PALAVRAS_IGNORADAS:
        if palavra in texto_lower:
            return True
    return False

def preencher_vazios_com_anteriores(relatorio_df, df_anteriores):
    if df_anteriores is None or df_anteriores.empty:
        return relatorio_df
    
    df_anteriores['Especialidade'] = df_anteriores['Especialidade'].replace('TO', 'Terapia Ocupacional')

    for index, row in relatorio_df.iterrows():
        if pd.isna(row['Anotacoes']) or str(row['Anotacoes']).strip() == '':
            especialidade = row['Especialidade']
            evolucoes_antigas_esp = df_anteriores[df_anteriores['Especialidade'] == especialidade]
            
            anotacao_substituta = ""
            for _, row_antiga in evolucoes_antigas_esp.iterrows():
                texto_antigo = str(row_antiga['Anotacoes_prontuario'])
                if not deve_ignorar_anotacao(texto_antigo):
                    anotacao_substituta = texto_antigo
                    break 
            
            if anotacao_substituta:
                relatorio_df.at[index, 'Anotacoes'] = anotacao_substituta
                
    return relatorio_df

def gerar_relatorio_docx(df_sessoes, df_evolucoes, df_evolucoes_anteriores, paciente):
    if paciente is None or df_sessoes is None:
        return None
    
    df_sessoes['Especialidade'] = df_sessoes['Especialidade'].replace('TO', 'Terapia Ocupacional')
    df_evolucoes['Especialidade'] = df_evolucoes['Especialidade'].replace('TO', 'Terapia Ocupacional')

    try:
        df_sessoes['Data_dt'] = pd.to_datetime(df_sessoes['Data'], format='%d/%m/%Y')
        df_sessoes = df_sessoes.sort_values(by='Data_dt', ascending=False).drop(columns=['Data_dt'])
    except ValueError:
         pass

    relatorio_df = pd.merge(df_sessoes, df_evolucoes, on=['Data', 'Especialidade'], how='left')
    mapa_paciente = PROFESSIONAL_MAP.get(paciente, {})
    relatorio_df['Profissional'] = relatorio_df['Especialidade'].map(mapa_paciente).fillna('A Definir')
    
    relatorio_df['Anotacoes'] = relatorio_df.get('Anotacoes_prontuario', '')
    relatorio_df['Anotacoes'] = relatorio_df['Anotacoes'].fillna('')

    relatorio_df = preencher_vazios_com_anteriores(relatorio_df, df_evolucoes_anteriores)

    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    document.add_heading(f'Relatório de Evolução - {paciente}', level=1)

    for _, row in relatorio_df.iterrows():
        if row['Data'] == 'Erro': continue

        p_prof = document.add_paragraph()
        p_prof.add_run('Profissional: ').bold = True
        p_prof.add_run(row['Profissional'])

        p_esp = document.add_paragraph()
        p_esp.add_run('Especialidade: ').bold = True
        p_esp.add_run(row['Especialidade'])

        p_data = document.add_paragraph()
        p_data.add_run('Data: ').bold = True
        p_data.add_run(row['Data'])
        
        p_anot = document.add_paragraph()
        p_anot.add_run('Anotações: ').bold = True
        p_anot.add_run(str(row['Anotacoes']))

        document.add_paragraph()

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def filtrar_sessoes_por_data(df_sessoes, data_inicio_str, data_fim_str):
    if not data_inicio_str or not data_fim_str:
        return df_sessoes
    try:
        df_copy = df_sessoes.copy()
        df_copy['Data_dt'] = pd.to_datetime(df_copy['Data'], format='%d/%m/%Y', errors='coerce')
        df_copy = df_copy.dropna(subset=['Data_dt'])

        data_inicio = pd.to_datetime(data_inicio_str, format='%Y-%m-%d')
        data_fim = pd.to_datetime(data_fim_str, format='%Y-%m-%d')
        
        df_filtrado = df_copy[(df_copy['Data_dt'] >= data_inicio) & (df_copy['Data_dt'] <= data_fim)].drop(columns=['Data_dt'])
        return df_filtrado
    except ValueError:
        return df_sessoes

# --- Rotas da Aplicação Web ---
@app.route('/', methods=['GET', 'POST'])
def index():
    error_message = None
    if request.method == 'POST':
        if 'controle_pdf' not in request.files or 'prontuario_pdf' not in request.files:
            return render_template('index.html', error="Erro: Faltando um ou mais arquivos.")

        controle_file = request.files['controle_pdf']
        prontuario_file = request.files['prontuario_pdf']
        prontuario_anterior_file = request.files.get('prontuario_anterior_pdf') 

        if controle_file.filename == '' or prontuario_file.filename == '':
            return render_template('index.html', error="Erro: Selecione os arquivos principais.")

        caminho_controle = os.path.join(app.config['UPLOAD_FOLDER'], controle_file.filename)
        caminho_prontuario = os.path.join(app.config['UPLOAD_FOLDER'], prontuario_file.filename)
        
        df_evolucoes_anteriores = None 
        caminho_anterior = None

        try:
            controle_file.save(caminho_controle)
            prontuario_file.save(caminho_prontuario)
            
            if prontuario_anterior_file and prontuario_anterior_file.filename != '':
                caminho_anterior = os.path.join(app.config['UPLOAD_FOLDER'], prontuario_anterior_file.filename)
                prontuario_anterior_file.save(caminho_anterior)
                texto_anterior = extrair_texto_de_pdf(caminho_anterior)
                df_evolucoes_anteriores = analisar_prontuario(texto_anterior)

            texto_controle = extrair_texto_de_pdf(caminho_controle)
            texto_prontuario = extrair_texto_de_pdf(caminho_prontuario)
            
            df_sessoes, paciente = get_schedule_for_patient(texto_controle)
            if paciente is None:
                raise ValueError("Paciente não reconhecido.")

            data_inicio = request.form.get('data_inicio')
            data_fim = request.form.get('data_fim')
            df_sessoes_filtradas = filtrar_sessoes_por_data(df_sessoes, data_inicio, data_fim)
            
            df_evolucoes = analisar_prontuario(texto_prontuario)
            
            file_stream = gerar_relatorio_docx(df_sessoes_filtradas, df_evolucoes, df_evolucoes_anteriores, paciente)

            primeiro_nome = paciente.split(' ')[0]
            nome_arquivo = f"Relatorio_{primeiro_nome}_{datetime.now().strftime('%Y-%m-%d')}.docx"

            return send_file(
                file_stream,
                as_attachment=True,
                download_name=nome_arquivo,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        
        except Exception as e:
            error_message = f"Erro ao processar arquivos: {e}"
            print(error_message)
            
        finally:
            if os.path.exists(caminho_controle): os.remove(caminho_controle)
            if os.path.exists(caminho_prontuario): os.remove(caminho_prontuario)
            if caminho_anterior and os.path.exists(caminho_anterior): os.remove(caminho_anterior)
                
    return render_template('index.html', error=error_message)

if __name__ == '__main__':
    app.run(debug=False)
